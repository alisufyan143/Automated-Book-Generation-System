"""
Document Compiler module for the Book Generation System.
Handles Stage 3: Final book compilation into .docx, .pdf, and .txt formats.
With improved formatting and professional styling.
"""

import os
import re
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

from database import Database
from config import Config


class MarkdownParser:
    """Parse and clean markdown content for document formatting."""
    
    @staticmethod
    def clean_markdown(text: str) -> str:
        """Remove raw markdown syntax and clean up text."""
        if not text:
            return ""
        
        # Remove ** bold markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # Remove * italic markers
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        # Remove __ bold markers
        text = re.sub(r'__([^_]+)__', r'\1', text)
        # Remove _ italic markers
        text = re.sub(r'_([^_]+)_', r'\1', text)
        # Remove ` code markers
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # Remove [link](url) - keep link text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Remove --- or *** horizontal rules
        text = re.sub(r'^[-*]{3,}$', '', text, flags=re.MULTILINE)
        # Clean up multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    @staticmethod
    def extract_bold_text(text: str) -> List[Tuple[str, bool]]:
        """Extract text with bold formatting info."""
        parts = []
        pattern = r'\*\*([^*]+)\*\*'
        last_end = 0
        
        for match in re.finditer(pattern, text):
            # Add text before the match
            if match.start() > last_end:
                parts.append((text[last_end:match.start()], False))
            # Add the bold text
            parts.append((match.group(1), True))
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            parts.append((text[last_end:], False))
        
        return parts if parts else [(text, False)]
    
    @staticmethod
    def parse_content(content: str) -> List[Dict]:
        """Parse content into structured blocks."""
        blocks = []
        
        # Split into paragraphs
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Detect block type
            if para.startswith('# '):
                blocks.append({'type': 'h1', 'text': para[2:].strip()})
            elif para.startswith('## '):
                blocks.append({'type': 'h2', 'text': para[3:].strip()})
            elif para.startswith('### '):
                blocks.append({'type': 'h3', 'text': para[4:].strip()})
            elif para.startswith('#### '):
                blocks.append({'type': 'h4', 'text': para[5:].strip()})
            elif para.startswith('- ') or para.startswith('* '):
                # Bullet list
                items = [line.strip()[2:] for line in para.split('\n') if line.strip().startswith(('- ', '* '))]
                blocks.append({'type': 'bullet_list', 'items': items})
            elif re.match(r'^\d+\.\s', para):
                # Numbered list
                items = []
                for line in para.split('\n'):
                    match = re.match(r'^\d+\.\s*(.+)', line.strip())
                    if match:
                        items.append(match.group(1))
                blocks.append({'type': 'numbered_list', 'items': items})
            elif para.startswith('> '):
                # Blockquote
                quote_text = '\n'.join(line[2:] if line.startswith('> ') else line for line in para.split('\n'))
                blocks.append({'type': 'quote', 'text': quote_text})
            elif para.startswith('```'):
                # Code block - just include as plain text
                code = para.strip('`').strip()
                if code.split('\n')[0].isalpha():  # Language identifier
                    code = '\n'.join(code.split('\n')[1:])
                blocks.append({'type': 'code', 'text': code})
            else:
                # Regular paragraph
                blocks.append({'type': 'paragraph', 'text': para})
        
        return blocks


class BookCompiler:
    """Handles final book compilation into various formats with professional formatting."""
    
    def __init__(self):
        """Initialize compiler with database connection."""
        self.db = Database()
        self.output_dir = Config.OUTPUT_DIR
        self.output_dir.mkdir(exist_ok=True)
        self.parser = MarkdownParser()
    
    # ==========================================================================
    # GATING LOGIC
    # ==========================================================================
    
    def check_compilation_ready(self, book_id: str) -> Dict[str, Any]:
        """Check if book is ready for final compilation."""
        book = self.db.get_book(book_id)
        if not book:
            return {"can_compile": False, "error": "Book not found"}
        
        chapters = self.db.get_book_chapters(book_id)
        issues = []
        
        final_status = book.get('final_review_notes_status', 'pending')
        if final_status == 'yes':
            issues.append("Waiting for final review notes")
        
        if not chapters:
            issues.append("No chapters found")
        else:
            pending = [c for c in chapters if not c.get('content')]
            if pending:
                issues.append(f"{len(pending)} chapters not generated")
        
        return {
            "can_compile": len(issues) == 0,
            "issues": issues,
            "total_chapters": len(chapters) if chapters else 0,
            "generated": len([c for c in chapters if c.get('content')]) if chapters else 0,
            "approved": len([c for c in chapters if c.get('status') == 'approved']) if chapters else 0,
            "final_status": final_status
        }
    
    # ==========================================================================
    # DOCX COMPILATION
    # ==========================================================================
    
    def _setup_docx_styles(self, doc: Document):
        """Configure document styles for professional appearance."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Georgia'
        font.size = Pt(11)
        
        # Heading 1 style
        h1_style = doc.styles['Heading 1']
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(24)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        
        # Heading 2 style
        h2_style = doc.styles['Heading 2']
        h2_style.font.name = 'Arial'
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        
        # Heading 3 style
        h3_style = doc.styles['Heading 3']
        h3_style.font.name = 'Arial'
        h3_style.font.size = Pt(13)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    
    def _add_formatted_paragraph(self, doc: Document, text: str, indent: bool = False):
        """Add a paragraph with proper formatting, handling markdown."""
        # Clean markdown but preserve structure
        clean_text = self.parser.clean_markdown(text)
        
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(10)
        para.paragraph_format.line_spacing = 1.5
        
        if indent:
            para.paragraph_format.left_indent = Inches(0.5)
        
        # Handle bold text within paragraph
        parts = self.parser.extract_bold_text(text)
        for text_part, is_bold in parts:
            clean_part = self.parser.clean_markdown(text_part)
            run = para.add_run(clean_part)
            if is_bold:
                run.bold = True
    
    def _add_bullet_list(self, doc: Document, items: List[str]):
        """Add a formatted bullet list."""
        for item in items:
            para = doc.add_paragraph(style='List Bullet')
            clean_item = self.parser.clean_markdown(item)
            para.add_run(clean_item)
            para.paragraph_format.space_after = Pt(4)
    
    def _add_numbered_list(self, doc: Document, items: List[str]):
        """Add a formatted numbered list."""
        for item in items:
            para = doc.add_paragraph(style='List Number')
            clean_item = self.parser.clean_markdown(item)
            para.add_run(clean_item)
            para.paragraph_format.space_after = Pt(4)
    
    def compile_to_docx(self, book_id: str, force: bool = False) -> str:
        """Compile book to professionally formatted Word document."""
        book = self.db.get_book(book_id)
        if not book:
            raise ValueError("Book not found")
        
        if not force:
            status = self.check_compilation_ready(book_id)
            if not status['can_compile']:
                raise ValueError(f"Cannot compile: {', '.join(status['issues'])}")
        
        chapters = self.db.get_book_chapters(book_id)
        chapters = [c for c in chapters if c.get('content')]
        
        # Create document
        doc = Document()
        self._setup_docx_styles(doc)
        
        # ===== TITLE PAGE =====
        for _ in range(6):
            doc.add_paragraph()
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(book['title'])
        title_run.bold = True
        title_run.font.size = Pt(36)
        title_run.font.name = 'Arial'
        title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        
        doc.add_paragraph()
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        
        doc.add_page_break()
        
        # ===== TABLE OF CONTENTS =====
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph()
        
        for chapter in chapters:
            toc_para = doc.add_paragraph()
            toc_para.paragraph_format.space_after = Pt(8)
            toc_run = toc_para.add_run(f"Chapter {chapter['chapter_number']}: {chapter.get('title', 'Untitled')}")
            toc_run.font.size = Pt(12)
        
        doc.add_page_break()
        
        # ===== CHAPTERS =====
        for chapter in chapters:
            # Chapter title
            doc.add_heading(
                f"Chapter {chapter['chapter_number']}: {chapter.get('title', 'Untitled')}", 
                level=1
            )
            
            # Parse and format content
            content = chapter.get('content', '')
            blocks = self.parser.parse_content(content)
            
            for block in blocks:
                if block['type'] == 'h1':
                    doc.add_heading(self.parser.clean_markdown(block['text']), level=1)
                elif block['type'] == 'h2':
                    doc.add_heading(self.parser.clean_markdown(block['text']), level=2)
                elif block['type'] == 'h3':
                    doc.add_heading(self.parser.clean_markdown(block['text']), level=3)
                elif block['type'] == 'h4':
                    h4 = doc.add_paragraph()
                    h4_run = h4.add_run(self.parser.clean_markdown(block['text']))
                    h4_run.bold = True
                    h4_run.font.size = Pt(12)
                elif block['type'] == 'bullet_list':
                    self._add_bullet_list(doc, block['items'])
                elif block['type'] == 'numbered_list':
                    self._add_numbered_list(doc, block['items'])
                elif block['type'] == 'quote':
                    quote_para = doc.add_paragraph()
                    quote_para.paragraph_format.left_indent = Inches(0.5)
                    quote_para.paragraph_format.right_indent = Inches(0.5)
                    quote_run = quote_para.add_run(self.parser.clean_markdown(block['text']))
                    quote_run.italic = True
                    quote_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
                elif block['type'] == 'code':
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.left_indent = Inches(0.25)
                    code_run = code_para.add_run(block['text'])
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                else:
                    self._add_formatted_paragraph(doc, block['text'])
            
            doc.add_page_break()
        
        # Save document
        safe_title = "".join(c for c in book['title'] if c.isalnum() or c in ' -_').strip()
        filename = f"{safe_title}.docx"
        filepath = self.output_dir / filename
        
        doc.save(filepath)
        self.db.update_book(book_id, output_docx_path=str(filepath))
        
        return str(filepath)
    
    # ==========================================================================
    # PDF COMPILATION
    # ==========================================================================
    
    def compile_to_pdf(self, book_id: str, force: bool = False) -> str:
        """Compile book to professionally formatted PDF."""
        book = self.db.get_book(book_id)
        if not book:
            raise ValueError("Book not found")
        
        if not force:
            status = self.check_compilation_ready(book_id)
            if not status['can_compile']:
                raise ValueError(f"Cannot compile: {', '.join(status['issues'])}")
        
        chapters = self.db.get_book_chapters(book_id)
        chapters = [c for c in chapters if c.get('content')]
        
        # Create PDF
        safe_title = "".join(c for c in book['title'] if c.isalnum() or c in ' -_').strip()
        filename = f"{safe_title}.pdf"
        filepath = self.output_dir / filename
        
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Import reportlab color
        from reportlab.lib.colors import HexColor
        
        # Styles
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'BookTitle',
            parent=styles['Heading1'],
            fontSize=32,
            alignment=TA_CENTER,
            spaceAfter=20,
            textColor=HexColor('#2C3E50')
        )
        
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceBefore=30,
            spaceAfter=20,
            textColor=HexColor('#2C3E50')
        )
        
        heading2_style = ParagraphStyle(
            'Heading2Custom',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=16,
            spaceAfter=10,
            textColor=HexColor('#34495E')
        )
        
        heading3_style = ParagraphStyle(
            'Heading3Custom',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=8,
            textColor=HexColor('#5D6D7E')
        )
        
        body_style = ParagraphStyle(
            'BookBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            firstLineIndent=24
        )
        
        quote_style = ParagraphStyle(
            'Quote',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            leftIndent=36,
            rightIndent=36,
            textColor='#7F8C8D',
            fontName='Times-Italic',
            spaceAfter=12
        )
        
        # Build story
        story = []
        
        # ===== TITLE PAGE =====
        story.append(Spacer(1, 3*inch))
        story.append(Paragraph(book['title'], title_style))
        story.append(Spacer(1, 0.5*inch))
        
        date_style = ParagraphStyle('Date', parent=styles['Normal'], alignment=TA_CENTER, textColor='#7F8C8D')
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", date_style))
        story.append(PageBreak())
        
        # ===== TABLE OF CONTENTS =====
        story.append(Paragraph("Table of Contents", chapter_title_style))
        story.append(Spacer(1, 0.3*inch))
        
        toc_style = ParagraphStyle('TOC', parent=styles['Normal'], fontSize=12, spaceAfter=8)
        for chapter in chapters:
            story.append(Paragraph(
                f"Chapter {chapter['chapter_number']}: {chapter.get('title', 'Untitled')}",
                toc_style
            ))
        story.append(PageBreak())
        
        # ===== CHAPTERS =====
        for chapter in chapters:
            story.append(Paragraph(
                f"Chapter {chapter['chapter_number']}: {chapter.get('title', 'Untitled')}",
                chapter_title_style
            ))
            
            # Parse content
            content = chapter.get('content', '')
            blocks = self.parser.parse_content(content)
            
            for block in blocks:
                text = self.parser.clean_markdown(block.get('text', ''))
                
                # Escape XML special characters
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                if block['type'] in ('h1', 'h2'):
                    story.append(Paragraph(text, heading2_style))
                elif block['type'] in ('h3', 'h4'):
                    story.append(Paragraph(text, heading3_style))
                elif block['type'] == 'bullet_list':
                    for item in block['items']:
                        clean_item = self.parser.clean_markdown(item)
                        clean_item = clean_item.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(f"• {clean_item}", body_style))
                elif block['type'] == 'numbered_list':
                    for i, item in enumerate(block['items'], 1):
                        clean_item = self.parser.clean_markdown(item)
                        clean_item = clean_item.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(f"{i}. {clean_item}", body_style))
                elif block['type'] == 'quote':
                    story.append(Paragraph(text, quote_style))
                else:
                    if text:
                        story.append(Paragraph(text, body_style))
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        self.db.update_book(book_id, output_pdf_path=str(filepath))
        
        return str(filepath)
    
    # ==========================================================================
    # TXT COMPILATION
    # ==========================================================================
    
    def compile_to_txt(self, book_id: str, force: bool = False) -> str:
        """Compile book to clean, readable plain text."""
        book = self.db.get_book(book_id)
        if not book:
            raise ValueError("Book not found")
        
        if not force:
            status = self.check_compilation_ready(book_id)
            if not status['can_compile']:
                raise ValueError(f"Cannot compile: {', '.join(status['issues'])}")
        
        chapters = self.db.get_book_chapters(book_id)
        chapters = [c for c in chapters if c.get('content')]
        
        lines = []
        width = 72
        
        # ===== TITLE PAGE =====
        lines.append("")
        lines.append("=" * width)
        lines.append("")
        lines.append(book['title'].center(width))
        lines.append("")
        lines.append("=" * width)
        lines.append("")
        lines.append(f"Generated on {datetime.now().strftime('%B %d, %Y')}".center(width))
        lines.append("")
        lines.append("")
        lines.append("")
        
        # ===== TABLE OF CONTENTS =====
        lines.append("-" * width)
        lines.append("TABLE OF CONTENTS".center(width))
        lines.append("-" * width)
        lines.append("")
        
        for chapter in chapters:
            lines.append(f"    Chapter {chapter['chapter_number']}: {chapter.get('title', 'Untitled')}")
        
        lines.append("")
        lines.append("")
        
        # ===== CHAPTERS =====
        for chapter in chapters:
            lines.append("=" * width)
            lines.append(f"CHAPTER {chapter['chapter_number']}")
            lines.append(chapter.get('title', 'Untitled').upper())
            lines.append("=" * width)
            lines.append("")
            
            # Parse and format content
            content = chapter.get('content', '')
            blocks = self.parser.parse_content(content)
            
            for block in blocks:
                text = self.parser.clean_markdown(block.get('text', ''))
                
                if block['type'] in ('h1', 'h2'):
                    lines.append("")
                    lines.append(text.upper())
                    lines.append("-" * len(text))
                    lines.append("")
                elif block['type'] in ('h3', 'h4'):
                    lines.append("")
                    lines.append(f"  {text}")
                    lines.append("")
                elif block['type'] == 'bullet_list':
                    for item in block['items']:
                        clean_item = self.parser.clean_markdown(item)
                        lines.append(f"    • {clean_item}")
                    lines.append("")
                elif block['type'] == 'numbered_list':
                    for i, item in enumerate(block['items'], 1):
                        clean_item = self.parser.clean_markdown(item)
                        lines.append(f"    {i}. {clean_item}")
                    lines.append("")
                elif block['type'] == 'quote':
                    lines.append(f'    "{text}"')
                    lines.append("")
                else:
                    # Word wrap paragraphs
                    if text:
                        wrapped = self._word_wrap(text, width - 4)
                        for line in wrapped:
                            lines.append(f"    {line}")
                        lines.append("")
            
            lines.append("")
            lines.append("")
        
        # Save file
        safe_title = "".join(c for c in book['title'] if c.isalnum() or c in ' -_').strip()
        filename = f"{safe_title}.txt"
        filepath = self.output_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        self.db.update_book(book_id, output_txt_path=str(filepath))
        
        return str(filepath)
    
    def _word_wrap(self, text: str, width: int) -> List[str]:
        """Word wrap text to specified width."""
        words = text.split()
        lines = []
        current_line = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 <= width:
                current_line.append(word)
                current_length += len(word) + 1
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return lines if lines else ['']
    
    # ==========================================================================
    # COMPILE ALL FORMATS
    # ==========================================================================
    
    def compile_book(self, book_id: str, formats: List[str] = None, force: bool = False) -> Dict[str, str]:
        """Compile book to all specified formats."""
        if formats is None:
            formats = ['docx', 'pdf', 'txt']
        
        results = {}
        
        for fmt in formats:
            print(f"📄 Compiling to {fmt.upper()}...")
            try:
                if fmt == 'docx':
                    results['docx'] = self.compile_to_docx(book_id, force)
                elif fmt == 'pdf':
                    results['pdf'] = self.compile_to_pdf(book_id, force)
                elif fmt == 'txt':
                    results['txt'] = self.compile_to_txt(book_id, force)
                print(f"   ✅ {results[fmt]}")
            except Exception as e:
                print(f"   ❌ Error: {e}")
                results[fmt] = f"Error: {e}"
        
        if all(not str(v).startswith('Error') for v in results.values()):
            self.db.update_book(book_id, book_output_status='completed')
            print("\n✅ Book compilation complete!")
        
        return results
    
    def approve_final_review(self, book_id: str) -> Dict[str, Any]:
        """Mark book as ready for final compilation."""
        self.db.update_book(book_id, final_review_notes_status='no_notes_needed')
        return {"success": True, "message": "Book approved for compilation"}
